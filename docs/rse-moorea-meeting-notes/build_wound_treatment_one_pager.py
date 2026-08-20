from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("coral-gardeners-wound-treatment-experiment-one-pager.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9E2E1", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(31, 77, 120)
    return p


def add_body(doc, text, after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(32, 43, 54)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(9.2)
    run.font.color.rgb = RGBColor(32, 43, 54)
    return p


def add_label_run(paragraph, label, text):
    r = paragraph.add_run(label)
    r.bold = True
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    r.font.size = Pt(9.2)
    r.font.color.rgb = RGBColor(32, 43, 54)
    r2 = paragraph.add_run(text)
    r2.font.name = "Calibri"
    r2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    r2.font.size = Pt(9.2)
    r2.font.color.rgb = RGBColor(32, 43, 54)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.68)
section.bottom_margin = Inches(0.62)
section.left_margin = Inches(0.72)
section.right_margin = Inches(0.72)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
styles["Normal"].font.size = Pt(9.5)
styles["List Bullet"].font.name = "Calibri"
styles["List Bullet"]._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
styles["List Bullet"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
styles["List Bullet"].font.size = Pt(9.2)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_after = Pt(1)
run = title.add_run("Small nursery test for coral wounds")
run.bold = True
run.font.name = "Calibri"
run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(22, 34, 45)

sub = doc.add_paragraph()
sub.paragraph_format.space_after = Pt(6)
sub_run = sub.add_run("Draft one-pager for Hannah Stewart and Coral Gardeners")
sub_run.font.name = "Calibri"
sub_run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
sub_run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
sub_run.font.size = Pt(9.5)
sub_run.font.color.rgb = RGBColor(85, 97, 110)

add_heading(doc, "Plain-language idea")
add_body(
    doc,
    "When coral fragments are made, some surfaces are left as fresh exposed skeleton. "
    "Those spots can foul, collect sediment, or heal slowly. We could run a small nursery test "
    "to ask whether covering those wounds helps enough to justify the extra time and cost.",
    after=4,
)

table = doc.add_table(rows=1, cols=4)
table.autofit = False
set_table_borders(table)
widths = [Inches(1.72), Inches(1.72), Inches(1.72), Inches(1.72)]
labels = [
    ("Open wound", "No cover. This is the comparison group."),
    ("Concrete", "Cheap mineral cover."),
    ("Apoxie Sculpt", "Epoxy benchmark Adrian has used."),
    ("Scripps material", "Only include if it can safely cover a fresh wound."),
]
fills = ["F7F9FB", "F7F9FB", "E8F3F9", "FFF5DF"]
for i, cell in enumerate(table.rows[0].cells):
    cell.width = widths[i]
    set_cell_margins(cell)
    set_cell_shading(cell, fills[i])
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    add_label_run(p, labels[i][0] + "\n", labels[i][1])

add_heading(doc, "Simple design")
add_bullet(doc, "Use healthy nursery fragments, ideally from known parent colonies.")
add_bullet(doc, "Make the same small wound on each fragment.")
add_bullet(doc, "Randomly assign fragments to the four treatments above.")
add_bullet(doc, "Keep each cover type spread across parent colonies, so we can tell whether the cover matters.")
add_bullet(doc, "Return fragments to the nursery and avoid placing one treatment only in one good or bad spot.")

add_heading(doc, "What we would measure")
add_bullet(doc, "Does the material stay on the wound?")
add_bullet(doc, "How fast does live tissue grow back over the wound?")
add_bullet(doc, "How much algae, sediment, or other buildup appears on the wound or cover?")
add_bullet(doc, "Is there any tissue loss, disease, bleaching, or death?")
add_bullet(doc, "Is there visible new skeleton or scar infill?")
add_bullet(doc, "How long does each treatment take to prepare and apply?")

add_heading(doc, "Suggested checks through time")
timeline = doc.add_table(rows=2, cols=5)
timeline.autofit = False
set_table_borders(timeline, color="D9E2E1", size="4")
times = ["T0", "2 weeks", "1 month", "3 months", "6 months"]
purposes = ["photo + wound size", "early loss + algae", "early closure", "main healing check", "cost/value decision"]
for i in range(5):
    for row, text, fill in ((0, times[i], "E8EEF5"), (1, purposes[i], "FFFFFF")):
        cell = timeline.cell(row, i)
        cell.width = Inches(1.38)
        set_cell_margins(cell, top=60, bottom=60, start=80, end=80)
        set_cell_shading(cell, fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        r.font.name = "Calibri"
        r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        r.font.size = Pt(8.5 if row else 9.0)
        r.bold = row == 0
        r.font.color.rgb = RGBColor(32, 43, 54)

add_heading(doc, "Important first decision")
add_body(
    doc,
    "Before treating the Scripps / Hybrid Reefs material as a wound cover, we should confirm what it is and what it is meant to do. "
    "If it is mainly a plug, tile, or baby-coral settlement material, it may be better as a separate nursery-surface test rather than a direct wound-cover test.",
    after=4,
)

add_heading(doc, "Questions for Hannah")
add_bullet(doc, "Which wound is most useful to test: a cut fragment end, or a side scrape like a donor-colony wound?")
add_bullet(doc, "What exactly is the Scripps / Hybrid Reefs material, and can it touch fresh exposed skeleton?")
add_bullet(doc, "Could Coral Gardeners spare a small set of nursery fragments from known parent colonies?")
add_bullet(doc, "Would a 2-week, 1-month, 3-month, and 6-month check fit the normal nursery workflow?")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fr = footer.add_run("RSE-Moorea / Coral Gardeners scoping draft")
fr.font.name = "Calibri"
fr._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
fr._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor(120, 130, 140)

doc.save(OUT)
print(OUT)
