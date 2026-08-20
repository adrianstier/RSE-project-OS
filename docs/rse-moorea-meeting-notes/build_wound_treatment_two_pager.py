from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
OUT = HERE / "coral-gardeners-wound-treatment-experiment-two-pager.docx"
FIGURE = HERE / "wound-treatment-experiment-figure.png"


def set_font(run, size=9.4, bold=False, color=(32, 43, 54)):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9E2E1", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_heading(doc, text, before=6, after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_font(run, size=10.5, bold=True, color=(31, 77, 120))
    return p


def add_body(doc, text, after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    set_font(run, size=9.4)
    return p


def add_bullet(doc, text, size=9.1, after=1.3):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_font(run, size=size)
    return p


def add_label_detail(cell, label, detail, detail_size=9.1):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(label + "\n")
    set_font(r1, size=9.2, bold=True)
    r2 = p.add_run(detail)
    set_font(r2, size=detail_size)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(title)
    set_font(run, size=18, bold=True, color=(22, 34, 45))
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    run2 = p2.add_run(subtitle)
    set_font(run2, size=9.5, color=(85, 97, 110))


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.62)
section.bottom_margin = Inches(0.55)
section.left_margin = Inches(0.65)
section.right_margin = Inches(0.65)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
styles["Normal"].font.size = Pt(9.4)
styles["List Bullet"].font.name = "Calibri"
styles["List Bullet"]._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
styles["List Bullet"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
styles["List Bullet"].font.size = Pt(9.1)

add_title(doc, "Nursery wound-cover test", "Short note for Hannah Stewart and Coral Gardeners")

add_heading(doc, "Why test this")
add_body(
    doc,
    "When coral fragments are cut, the new edge leaves exposed skeleton. Those spots can collect algae, "
    "sediment, and other fouling before tissue grows back. A small nursery test can tell us whether covering "
    "the wound improves healing enough to justify the added handling time.",
)

add_heading(doc, "Cover types to compare")
treatments = doc.add_table(rows=1, cols=4)
treatments.autofit = False
set_table_borders(treatments)
labels = [
    ("Open wound", "No cover. Baseline comparison."),
    ("Concrete", "Low-cost mineral cover."),
    ("Apoxie Sculpt", "Epoxy treatment Adrian has used before."),
    ("Scripps material", "Confirm its use case, then include it as the new-material comparison."),
]
fills = ["F7F9FB", "F7F9FB", "E8F3F9", "FFF5DF"]
for i, cell in enumerate(treatments.rows[0].cells):
    cell.width = Inches(1.75)
    set_cell_margins(cell, top=80, bottom=80, start=100, end=100)
    set_cell_shading(cell, fills[i])
    add_label_detail(cell, labels[i][0], labels[i][1], detail_size=8.8)

add_heading(doc, "Design")
for item in [
    "Use healthy nursery fragments from known parent colonies where possible.",
    "Make one small, repeatable wound on each fragment.",
    "Assign fragments at random to open wound, concrete, Apoxie Sculpt, or the Scripps material.",
    "Spread each cover type across parent colonies or genotypes, so the material effect is not mixed up with lineage.",
    "Spread treatments across the nursery instead of putting one cover type in one part of the array.",
    "Photograph each fragment the same way at every check.",
]:
    add_bullet(doc, item)

add_heading(doc, "What to measure")
for item in [
    "Cover retention: is the material still attached?",
    "Tissue regrowth: how much live tissue has grown back across the wound?",
    "Fouling: algae, sediment, and other buildup on the wound or cover.",
    "Coral health: tissue loss, disease signs, bleaching, or death.",
    "New skeleton: visible scar filling or new skeleton on the cut surface.",
    "Handling cost: minutes per fragment and any added material cost.",
]:
    add_bullet(doc, item)

add_heading(doc, "Checks")
timeline = doc.add_table(rows=2, cols=5)
timeline.autofit = False
set_table_borders(timeline, color="D9E2E1", size="4")
times = ["T0", "2 weeks", "1 month", "3 months", "6 months"]
purposes = ["photo + wound area", "attachment + fouling", "early healing", "healing + health", "cost/value call"]
for i in range(5):
    for row, text, fill in ((0, times[i], "E8EEF5"), (1, purposes[i], "FFFFFF")):
        cell = timeline.cell(row, i)
        cell.width = Inches(1.4)
        set_cell_margins(cell, top=60, bottom=60, start=80, end=80)
        set_cell_shading(cell, fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_font(run, size=8.5 if row else 9.0, bold=(row == 0))

add_heading(doc, "Scripps material check")
add_body(
    doc,
    "The Scripps / Hybrid Reefs material belongs in this test if it can sit on a fresh coral wound. "
    "First, ask what the material is, how it cures or breaks down, whether it has touched living coral tissue before, "
    "and whether it is a cover, plug, tile, or settlement surface. If it is mainly a settlement surface, keep it as a separate nursery-surface test.",
    after=2,
)

doc.add_page_break()

add_heading(doc, "Experiment design", before=0, after=4)
fig_p = doc.add_paragraph()
fig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fig_p.paragraph_format.space_after = Pt(4)
run = fig_p.add_run()
inline = run.add_picture(str(FIGURE), width=Inches(7.05))
inline._inline.docPr.set("descr", "Experiment design figure showing material triage, randomized nursery-fragment wound treatments, monitoring, and decision rule.")

cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap.paragraph_format.space_after = Pt(7)
cap_run = cap.add_run(
    "Figure. The design keeps the test small, uses Coral Gardeners nursery fragments, "
    "and compares wound covers against an open-wound baseline."
)
set_font(cap_run, size=8.6, color=(85, 97, 110))

add_heading(doc, "Figure notes", before=2, after=2)
for item in [
    "Panel A: check the Scripps / Hybrid Reefs material first: composition, cure or breakdown, and contact with living coral.",
    "Panel B: make one repeatable wound per fragment, assign treatments at random, and track the same measures through time.",
    "Panel C: keep a cover only if it heals faster, does not create health problems, and is fast enough for nursery work.",
]:
    add_bullet(doc, item, size=8.9, after=1.0)

add_heading(doc, "Questions for Hannah", before=4, after=2)
for item in [
    "Which wound matters most for Coral Gardeners: a cut fragment end, or a side scrape like a donor-colony wound?",
    "What is the Scripps / Hybrid Reefs material: cover, plug, tile, or settlement surface?",
    "How many nursery fragments from known parent colonies can be used without disrupting production?",
    "Do the 2-week, 1-month, 3-month, and 6-month checks fit the normal nursery schedule?",
]:
    add_bullet(doc, item, size=8.9, after=0.8)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fr = footer.add_run("RSE-Moorea / Coral Gardeners scoping draft")
set_font(fr, size=8, color=(120, 130, 140))

doc.save(OUT)
print(OUT)
